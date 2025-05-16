from fastapi import FastAPI, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse
from pydantic import BaseModel
from pymongo import MongoClient
from typing import Optional, Dict, List
from reportlab.lib.pagesizes import letter, landscape
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib import colors
import bcrypt
import pandas as pd
import openpyxl
from io import BytesIO
from datetime import datetime
from fastapi import Depends
from bson import ObjectId
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from docx import Document
from docx.shared import Pt
from fastapi import Query



app = FastAPI()

# CORS
origins = [
    "https://cateringbestem.onrender.com",
    "https://catering-1.onrender.com",
    "http://localhost",
    "http://127.0.0.1",
    "http://localhost:5500",
    "null"
]
app.add_middleware(
    CORSMiddleware,
    allow_origins=origins,
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# MongoDB
client = MongoClient("mongodb+srv://Maciej:20250504@cateringAtlas.ivdqqew.mongodb.net/catering_app?retryWrites=true&w=majority")
db = client["catering_app"]
users_collection = db["users"]
orders_collection = db["orders"]
menu_collection = db["menu"]
messages_collection = db["messages"]

# MODELE
class LoginUser(BaseModel):
    username: str
    password: str

class Meal(BaseModel):
    name: str
    price: float

class WeeklyOrder(BaseModel):
    username: str
    meals: Dict[str, List[Meal]]  # Klucz to dzieĹ„ tygodnia, wartoĹ›Ä‡ to lista obiektĂłw Meal
    week: str
    date_range: str


class MenuItem(BaseModel):
    name: str
    description: str
    price: float
    day: str  # Dodane pole dnia
    username: str

class NewUserPayload(BaseModel):
    username: str
    password: str
    role: str
    admin_username: str
    user_code: str

class MenuPayload(BaseModel):
    name: str
    description: str
    price: float
    username: str
    day: str

class MenuDeletePayload(BaseModel):
    name: str
    username: str

class DeleteOrderPayload(BaseModel):
    order_id: str
    admin_username: str

class DeleteOrdersPayload(BaseModel):
    order_ids: List[str]
    admin_username: str


class Message(BaseModel):
    text: str


# Funkcje haseĹ‚
def hash_password(password: str) -> str:
    return bcrypt.hashpw(password.encode('utf-8'), bcrypt.gensalt()).decode('utf-8')

def verify_password(password: str, hashed: str) -> bool:
    return bcrypt.checkpw(password.encode('utf-8'), hashed.encode('utf-8'))

# Endpointy

@app.post("/login")
def login(user: LoginUser):
    db_user = users_collection.find_one({"username": user.username})
    if not db_user or not verify_password(user.password, db_user["hashed_password"]):
        raise HTTPException(status_code=400, detail="Błędny login lub hasło")

    return {
        "msg": "Zalogowano",
        "username": db_user["username"],
        "role": db_user["role"],
        "user_code": db_user.get("user_code", "")  # Dodajemy user_code do odpowiedzi
    }

@app.post("/admin/add_user")
def add_user(payload: NewUserPayload):
    print(payload)  # SprawdĹş, co przychodzi w ĹĽÄ…daniu
    admin = users_collection.find_one({"username": payload.admin_username})
    if not admin or admin["role"] != "admin":
        raise HTTPException(status_code=403, detail="Brak dostępu")

    if users_collection.find_one({"username": payload.username}):
        raise HTTPException(status_code=400, detail="Użytkownik już istnieje")

    hashed_pw = hash_password(payload.password)
    users_collection.insert_one({
        "username": payload.username,
        "hashed_password": hashed_pw,
        "role": payload.role,
        "user_code": payload.user_code
    })
    return {"msg": f"użytkownik {payload.username} dodany jako {payload.role}"}


@app.post("/menu")
def add_menu_item(payload: MenuPayload):
    user = users_collection.find_one({"username": payload.username})
    if not user or user["role"] != "admin":
        raise HTTPException(status_code=403, detail="Brak uprawnień")

    menu_collection.insert_one({
        "name": payload.name,
        "description": payload.description,
        "price": payload.price,
        "day": payload.day  # Dodane pole dnia
    })
    return {"msg": "Pozycja dodana do menu"}

# Zmodyfikuj endpoint /menu/list


@app.get("/menu/list")
def get_menu():
    return list(menu_collection.find({}, {"_id": 0}))

@app.delete("/menu/delete")
def delete_menu_item(payload: MenuDeletePayload):
    admin = users_collection.find_one({"username": payload.username})
    if not admin or admin["role"] != "admin":
        raise HTTPException(status_code=403, detail="Brak uprawnień„")

    result = menu_collection.delete_one({"name": payload.name})
    if result.deleted_count == 0:
        raise HTTPException(status_code=404, detail="Pozycja nie istnieje")
    return {"msg": f"Usunięto pozycje™ {payload.name}"}

@app.delete("/admin/delete_user")
def delete_user(username: str, admin_username: str):
    admin = users_collection.find_one({"username": admin_username})
    if not admin or admin["role"] != "admin":
        raise HTTPException(status_code=403, detail="Brak dostępu")

    user = users_collection.find_one({"username": username})
    if not user:
        raise HTTPException(status_code=404, detail="Użytkownik nie istnieje")

    users_collection.delete_one({"username": username})

    return {"msg": f"Użytkownik {username} został‚ usunięty"}

@app.post("/order/weekly")
def create_weekly_order(order: WeeklyOrder):
    user = users_collection.find_one({"username": order.username})
    if not user:
        raise HTTPException(status_code=400, detail="Użytkownik nie istnieje")

    # ZamĂłwienie do zapisania
    order_data = order.dict()
    meals_data = []
    for day, meal_list in order.meals.items():
        for meal in meal_list:
            meals_data.append({
                "day": day,
                "name": meal.name,
                "price": meal.price
            })

    orders_collection.insert_one({
        "username": order.username,
        "meals": meals_data,
        "week": order.week,
        "date_range": order.date_range,
        "timestamp": datetime.utcnow()
    })
    return {"msg": "Zamówienie zapisane"}


from bson import ObjectId
from fastapi import HTTPException

from bson import ObjectId
from fastapi import HTTPException


@app.delete("/admin/delete_order")
async def delete_order(payload: DeleteOrderPayload):
    # Walidacja admina
    admin = users_collection.find_one({"username": payload.admin_username})
    if not admin or admin["role"] != "admin":
        raise HTTPException(status_code=403, detail="Brak uprawnień")

    # Konwersja stringa na ObjectId MongoDB
    try:
        object_id = ObjectId(payload.order_id)
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Nieprawidłowy format ID: {str(e)}")

    # Usuwanie po _id (ObjectId)
    result = orders_collection.delete_one({"_id": object_id})

    if result.deleted_count == 0:
        raise HTTPException(status_code=404, detail="Zamówienie nie istnieje")

    return {"status": "success", "message": f"Zamówienie {payload.order_id} usunięte"}

@app.get("/order/history")
def get_user_orders(username: str):
    user = users_collection.find_one({"username": username})
    if not user:
        raise HTTPException(status_code=403, detail="Brak użytkownika")
    return list(orders_collection.find({"username": username}, {"_id": 0}))

@app.put("/admin/change_password")
def change_password(username: str, new_password: str, admin_username: str):
    admin = users_collection.find_one({"username": admin_username})
    if not admin or admin["role"] != "admin":
        raise HTTPException(status_code=403, detail="Brak dostępu")

    user = users_collection.find_one({"username": username})
    if not user:
        raise HTTPException(status_code=404, detail="Użytkownik nie istnieje")

    hashed_password = hash_password(new_password)
    users_collection.update_one({"username": username}, {"$set": {"hashed_password": hashed_password}})

    return {"msg": f"Hasło użytkownika {username} zostało zmienione."}

@app.get("/admin/users")
def get_users(admin_username: str):
    admin = users_collection.find_one({"username": admin_username})
    if not admin or admin["role"] != "admin":
        raise HTTPException(status_code=403, detail="Brak dostępu")

    users = list(users_collection.find({}, {"_id": 0, "hashed_password": 0}))
    return users

@app.put("/admin/update_role")
def update_role(username: str, new_role: str, admin_username: str):
    admin = users_collection.find_one({"username": admin_username})
    if not admin or admin["role"] != "admin":
        raise HTTPException(status_code=403, detail="Brak dostępu")

    user = users_collection.find_one({"username": username})
    if not user:
        raise HTTPException(status_code=404, detail="Użytkownik nie istnieje")

    if new_role not in ["user", "admin"]:
        raise HTTPException(status_code=400, detail="Nieprawidłoowa rola")

    users_collection.update_one({"username": username}, {"$set": {"role": new_role}})

    return {"msg": f"Rola użytkownika {username} została zmieniona na {new_role}"}

# Nowy endpoint do pobrania raportu zamĂłwieĹ„ w formacie Excel
@app.get("/admin/orders/excel")
def export_orders_excel(admin_username: str):
    # Walidacja uprawnieĹ„ administratora
    admin = users_collection.find_one({"username": admin_username})
    if not admin or admin["role"] != "admin":
        raise HTTPException(status_code=403, detail="Brak uprawnień administratora")

    # Pobierz zamĂłwienia
    orders = list(orders_collection.find({}))
    if not orders:
        raise HTTPException(status_code=404, detail="Brak zamówień do eksportu")

    # Przygotuj dane
    rows = []
    days_of_week = ["Poniedziałek", "Wtorek", "Środa", "Czwartek", "Piątek"]

    for order in orders:
        # Pobierz kod uĹĽytkownika
        user = users_collection.find_one({"username": order["username"]})
        user_code = user.get("user_code", "") if user else ""

        # Inicjalizuj struktury dla dan i cen
        meals_by_day = {day: [] for day in days_of_week}
        prices_by_day = {day: 0.0 for day in days_of_week}

        # Grupuj dania i sumuj ceny wedĹ‚ug dnia
        for meal in order.get("meals", []):
            day = meal["day"]
            if day in meals_by_day:
                meals_by_day[day].append(meal['name'])
                prices_by_day[day] += float(meal['price'])

        # Przygotuj wiersz danych
        row = {
            "ID zamówienia": str(order["_id"]),
            "Kod użytkownika": user_code,
            "Użytkownik": order["username"],
            "Miejsce": order.get("date_range", ""),
            "Tydzień„": order["week"],
            "Data zamówienia": order.get("timestamp", "").strftime("%Y-%m-%d %H:%M:%S") if order.get(
                "timestamp") else "",

        }

        # Dodaj kolumny dla kaĹĽdego dnia (dania i ceny)
        for day in days_of_week:
            # Kolumna z daniami
            row[f"{day} - danie"] = ", ".join(meals_by_day[day]) if meals_by_day[day] else "Brak"
            # Kolumna z cenÄ… (tylko wartoĹ›Ä‡ liczbowa)
            row[f"{day} - cena"] = prices_by_day[day] if prices_by_day[day] > 0 else 0.0

        rows.append(row)

    # UtwĂłrz DataFrame
    df = pd.DataFrame(rows)

    # UtwĂłrz plik Excel w pamiÄ™ci
    output = BytesIO()
    with pd.ExcelWriter(output, engine='openpyxl') as writer:
        df.to_excel(writer, index=False, sheet_name='Zamówienia')
        worksheet = writer.sheets['Zamówienia']

        # Formatowanie kolumn z cenami jako waluta
        for col_idx, column in enumerate(df.columns, 1):
            if "cena" in column:
                # Ustaw formatowanie walutowe dla kolumn z cenami
                for row in range(2, len(df) + 2):
                    worksheet.cell(row=row, column=col_idx).number_format = '#,##0.00" zł‚"'

            # Dostosuj szerokoĹ›Ä‡ kolumn
            column_width = max(df[column].astype(str).map(len).max(), len(column)) + 2
            worksheet.column_dimensions[openpyxl.utils.get_column_letter(col_idx)].width = column_width

    output.seek(0)

    # Ustaw nagĹ‚Ăłwki odpowiedzi
    headers = {
        "Content-Disposition": "attachment; filename=raport_zamowien.xlsx",
        "Access-Control-Expose-Headers": "Content-Disposition"
    }

    return StreamingResponse(
        output,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers=headers
    )

@app.get("/admin/orders")
def get_all_orders(admin_username: str):
    admin = users_collection.find_one({"username": admin_username})
    if not admin or admin["role"] != "admin":
        raise HTTPException(status_code=403, detail="Brak uprawnień")

    try:
        orders = list(orders_collection.find({}))

        # Konwersja ObjectId na string dla kaĹĽdego zamĂłwienia
        for order in orders:
            order["_id"] = str(order["_id"])

        return orders

    except Exception as e:
        print(f"Błąd w /admin/orders: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Błąd serwera: {str(e)}")

@app.delete("/admin/delete_orders")
async def delete_orders(payload: DeleteOrdersPayload):
    # Walidacja admina
    admin = users_collection.find_one({"username": payload.admin_username})
    if not admin or admin["role"] != "admin":
        raise HTTPException(status_code=403, detail="Brak uprawnieĹ„")

    # Konwersja stringĂłw na ObjectId
    try:
        object_ids = [ObjectId(order_id) for order_id in payload.order_ids]
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Nieprawidłowy format ID: {str(e)}")

    # Usuwanie wielu zamĂłwieĹ„
    result = orders_collection.delete_many({"_id": {"$in": object_ids}})

    return {
        "status": "success",
        "message": f"UsuniÄ™to {result.deleted_count} zamĂłwieĹ„",
        "deleted_count": result.deleted_count
    }


@app.get("/admin/orders/pdf")
def export_orders_pdf(admin_username: str):
    # Walidacja uprawnieĹ„ administratora
    admin = users_collection.find_one({"username": admin_username})
    if not admin or admin["role"] != "admin":
        raise HTTPException(status_code=403, detail="Brak uprawnień administratora")

    try:
        # Pobierz zamĂłwienia
        orders = list(orders_collection.find({}))
        if not orders:
            raise HTTPException(status_code=404, detail="Brak zamĂłwień do eksportu")

        # UtwĂłrz PDF w pamiÄ™ci z orientacjÄ… poziomÄ…
        buffer = BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=landscape(letter))

        # Zarejestruj czcionkÄ™ (waĹĽne dla polskich znakĂłw)
        try:
            pdfmetrics.registerFont(TTFont('DejaVu', 'DejaVuSans.ttf'))
            styles = getSampleStyleSheet()
            styles['Normal'].fontName = 'DejaVu'
        except:
            print("Uwaga: Czcionka DejaVu nie została zarejestrowana, używam domyślnej")

        # Inicjalizacja elementĂłw dokumentu
        elements = []
        styles = getSampleStyleSheet()

        # TytuĹ‚
        title = Paragraph(f"Raport zamówień„ - {datetime.now().strftime('%Y-%m-%d')}", styles['Title'])
        elements.append(title)

        # Przygotuj dane
        data = []
        headers = [
            "Nazwisko Imie",
            "RCP",
            "Tydzień„",
            "Zakres dat",
            "Poniedziałek",
            "Wtorek",
            "Środa",
            "Czwartek",
            "Piątek"
        ]
        data.append(headers)

        for order in orders:
            # Pobierz kod uĹĽytkownika
            user = users_collection.find_one({"username": order["username"]})
            user_code = user.get("user_code", "") if user else ""

            # Grupuj dania wedĹ‚ug dnia
            meals_by_day = {day: [] for day in headers[4:]}
            for meal in order.get("meals", []):
                day = meal["day"]
                if day in meals_by_day:
                    meals_by_day[day].append(f"{meal['name']} ({meal['price']:.2f} zĹ‚)")

            # Przygotuj wiersz danych
            row = [
                user_code,
                order["username"],
                order["week"],
                order.get("date_range", "Brak danych"),
                "\n".join(meals_by_day["Poniedziałek"]) or "-",
                "\n".join(meals_by_day["Wtorek"]) or "-",
                "\n".join(meals_by_day["Środa"]) or "-",
                "\n".join(meals_by_day["Czwartek"]) or "-",
                "\n".join(meals_by_day["Piątek"]) or "-"
            ]
            data.append(row)

        # SzerokoĹ›ci kolumn
        col_widths = [80, 60, 50, 80] + [70] * 5  # Suma 9 kolumn

        # Tabela
        table = Table(data, colWidths=col_widths, repeatRows=1)

        # Style tabeli
        style = TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#3a86ff')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 8),
            ('FONTSIZE', (0, 1), (-1, -1), 7),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 8),
            ('BACKGROUND', (0, 1), (-1, -1), colors.white),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.lightgrey),
            ('VALIGN', (0, 0), (-1, -1), 'TOP')
        ])

        # Alternatywne kolory wierszy
        for i in range(1, len(data)):
            if i % 2 == 0:
                style.add('BACKGROUND', (0, i), (-1, i), colors.HexColor('#f8f9fa'))

        table.setStyle(style)
        elements.append(table)

        # Zbuduj dokument
        doc.build(elements)
        buffer.seek(0)

        # Ustaw nagĹ‚Ăłwki odpowiedzi
        headers = {
            "Content-Disposition": f"attachment; filename=raport_zamowien_{datetime.now().strftime('%Y-%m-%d')}.pdf",
            "Access-Control-Expose-Headers": "Content-Disposition"
        }

        return StreamingResponse(
            buffer,
            media_type="application/pdf",
            headers=headers
        )

    except Exception as e:
        print(f"Błąd podczas generowania PDF: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Błąd generowania PDF: {str(e)}")

@app.get("/admin/orders/erp")
def export_orders_erp(admin_username: str):
    # Walidacja uprawnieĹ„ administratora
    admin = users_collection.find_one({"username": admin_username})
    if not admin or admin["role"] != "admin":
        raise HTTPException(status_code=403, detail="Brak uprawnień administratora")

    # Pobierz zamĂłwienia
    orders = list(orders_collection.find({}))
    if not orders:
        raise HTTPException(status_code=404, detail="Brak zamówień do eksportu")

    # Przygotuj dane w formacie CSV
    output = BytesIO()

    # NagĹ‚Ăłwki kolumn
    headers = "username;total_price;week"
    output.write(headers.encode('utf-8'))

    for order in orders:
        # Oblicz sumÄ™ cen (zaokrÄ…glonÄ… do 2 miejsc po przecinku)
        total_price = round(sum(meal['price'] for meal in order.get("meals", [])), 2)

        # Przygotuj wiersz danych - total_price jako liczba bez dodatkowych oznaczeĹ„ i .0 dla liczb caĹ‚kowitych
        if total_price % 1 == 0:
            total_price = int(total_price)  # UsuĹ„ .0 dla liczb caĹ‚kowitych
        row = f"\n{order['username']};{total_price};{order['week']}"
        output.write(row.encode('utf-8'))

    output.seek(0)

    # Ustaw nagĹ‚Ăłwki odpowiedzi
    headers = {
        "Content-Disposition": "attachment; filename=raport_erp.csv",
        "Access-Control-Expose-Headers": "Content-Disposition",
        "Content-Type": "text/csv; charset=utf-8"
    }

    return StreamingResponse(
        output,
        media_type="text/csv",
        headers=headers
    )


@app.get("/admin/orders/dishes_report")
def export_dishes_report(admin_username: str):
    # Walidacja uprawnieĹ„ administratora
    admin = users_collection.find_one({"username": admin_username})
    if not admin or admin["role"] != "admin":
        raise HTTPException(status_code=403, detail="Brak uprawnień administratora")

    # Pobierz zamĂłwienia
    orders = list(orders_collection.find({}))
    if not orders:
        raise HTTPException(status_code=404, detail="Brak zamówień do eksportu")

    # Przygotuj dane
    rows = []

    for order in orders:
        # Pobierz kod uĹĽytkownika
        user = users_collection.find_one({"username": order["username"]})
        user_code = user.get("user_code", "") if user else ""

        # Dla kaĹĽdego dania w zamĂłwieniu utwĂłrz osobny wiersz
        for meal in order.get("meals", []):
            row = {
                "Kod użytkownika": user_code,
                "Miejsce": order.get("date_range", "Brak danych"),
                "Tydzień": order["week"],
                "Danie": f"{meal['name']} ({meal['day']})",
                "Cena": float(meal['price'])
            }
            rows.append(row)

    # UtwĂłrz DataFrame
    df = pd.DataFrame(rows)

    # UtwĂłrz plik Excel w pamiÄ™ci
    output = BytesIO()
    with pd.ExcelWriter(output, engine='openpyxl') as writer:
        df.to_excel(writer, index=False, sheet_name='Raport daĹ„')
        worksheet = writer.sheets['Raport daĹ„']

        # Formatowanie kolumny z cenÄ… jako waluta
        for row in range(2, len(df) + 2):
            worksheet.cell(row=row, column=5).number_format = '#,##0.00" zł"'

        # Dostosuj szerokoĹ›ci kolumn
        col_widths = {
            'A': 20,  # Kod uĹĽytkownika
            'B': 30,  # Miejsce
            'C': 15,  # TydzieĹ„
            'D': 50,  # Danie
            'E': 15  # Cena
        }

        for col, width in col_widths.items():
            worksheet.column_dimensions[col].width = width

    output.seek(0)

    # Ustaw nagĹ‚Ăłwki odpowiedzi
    headers = {
        "Content-Disposition": "attachment; filename=raport_dan.xlsx",
        "Access-Control-Expose-Headers": "Content-Disposition"
    }

    return StreamingResponse(
        output,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers=headers
    )


from docx import Document
from docx.shared import Pt


@app.get("/admin/orders/word_mailmerge")
def export_orders_word_mailmerge(admin_username: str):
    # Walidacja uprawnieĹ„ administratora
    admin = users_collection.find_one({"username": admin_username})
    if not admin or admin["role"] != "admin":
        raise HTTPException(status_code=403, detail="Brak uprawnieĹ„ administratora")

    # Pobierz zamĂłwienia
    orders = list(orders_collection.find({}))
    if not orders:
        raise HTTPException(status_code=404, detail="Brak zamówień do eksportu")

    # UtwĂłrz nowy dokument Word
    doc = Document()

    # NagĹ‚Ăłwek
    doc.add_heading('Raport zamówień - korespondencja seryjna', 0)

    # Tabela z danymi
    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'

    # NagĹ‚Ăłwki kolumn
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Kod użytkownika'
    hdr_cells[1].text = 'Nazwa użytkownika'
    hdr_cells[2].text = 'Tydzień'
    hdr_cells[3].text = 'Miejsce'
    hdr_cells[4].text = 'Dania'

    # WypeĹ‚nij tabelÄ™ danymi
    for order in orders:
        user = users_collection.find_one({"username": order["username"]})
        user_code = user.get("user_code", "") if user else ""

        meals = "\n".join([f"{m['day']}: {m['name']} ({m['price']} zł‚)" for m in order.get("meals", [])])

        row_cells = table.add_row().cells
        row_cells[0].text = user_code
        row_cells[1].text = order["username"]
        row_cells[2].text = order["week"]
        row_cells[3].text = order.get("date_range", "Brak danych")
        row_cells[4].text = meals

    # Zapisz do bufora
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    # Ustaw nagĹ‚Ăłwki odpowiedzi
    headers = {
        "Content-Disposition": "attachment; filename=raport_korespondencja.docx",
        "Access-Control-Expose-Headers": "Content-Disposition"
    }

    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers=headers
    )
@app.get("/order/exists")
def order_exists(username: str = Query(...), week: str = Query(...)):
    existing_order = orders_collection.find_one({"username": username, "week": week})
    return {"exists": existing_order is not None}

@app.put("/messages")
async def update_message(message: Message):
    result = messages_collection.update_one(
        {"_id": "login_info"},
        {"$set": {"text": message.text}},
        upsert=True
    )
    return {"status": "ok"}


@app.post("/messages")
async def add_message(message: Message):
    result = messages_collection.update_one(
        {"_id": "login_info"},
        {"$set": {"text": message.text}},
        upsert=True
    )
    return {"status": "ok"}

@app.get("/messages")
async def get_message():
    msg = messages_collection.find_one({"_id": "login_info"})
    return {"text": msg["text"] if msg else ""}
